import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def create_paragraph_style():
    style = {
        "first_line_indent": Cm(0.74),
        "line_spacing": 1,
        "alignment": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "font_size": Pt(10.5),
        "font_name_ascii": "Cambria",  # 西文字体
        "font_name_east_asia": "等线",  # 中文字体
    }
    return style


def set_document_style(doc, translated_paragraphs):
    # 设置页面边距
    for section in doc.sections:
        for margin in ["left_margin", "right_margin"]:
            setattr(section, margin, Cm(3.18))
        for margin in ["top_margin", "bottom_margin"]:
            setattr(section, margin, Cm(2.54))

    style = create_paragraph_style()

    for text in translated_paragraphs:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = style["line_spacing"]
        p.paragraph_format.first_line_indent = style["first_line_indent"]
        p.alignment = style["alignment"]

        run = p.add_run(text)
        run.font.size = style["font_size"]

        # 分别设置中英文字体
        run.font.name = style["font_name_ascii"]  # 西文字体
        run._element.rPr.rFonts.set(
            qn("w:ascii"), style["font_name_ascii"]
        )  # ASCII字符
        run._element.rPr.rFonts.set(
            qn("w:hAnsi"), style["font_name_ascii"]
        )  # 高位ASCII字符
        run._element.rPr.rFonts.set(
            qn("w:eastAsia"), style["font_name_east_asia"]
        )  # 中文字符

        # 禁用中文断行时的字符间距调整
        p.paragraph_format.word_wrap = True
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    return doc


def create_paragraph_properties(first_line_indent=0.74):
    """创建段落属性元素"""
    pPr = OxmlElement("w:pPr")

    # 添加缩进设置
    ind = OxmlElement("w:ind")
    ind.set(qn("w:firstLine"), str(int(first_line_indent * 567)))

    # 添加段落间距和行距设置
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "100")  # 段后5磅
    spacing.set(qn("w:line"), "240")  # 单倍行距
    spacing.set(qn("w:lineRule"), "auto")  # 行距规则为自动

    pPr.append(ind)
    pPr.append(spacing)
    return pPr


def create_new_paragraph(first_line_indent=0.74):
    """创建新段落元素"""
    new_p = OxmlElement("w:p")
    pPr = create_paragraph_properties(first_line_indent)
    new_p.append(pPr)
    return new_p


def handle_consecutive_breaks(br_elements, first_line_indent):
    """处理连续的换行符"""
    for i in range(len(br_elements) - 1):
        current_br = br_elements[i]
        next_br = br_elements[i + 1]

        if current_br.getnext() == next_br:
            new_p = create_new_paragraph(first_line_indent)

            # 在当前换行符位置插入新段落
            current_br.getparent().insert(
                current_br.getparent().index(current_br), new_p
            )
            # 移除连续的两个换行符
            current_br.getparent().remove(current_br)
            next_br.getparent().remove(next_br)


def apply_paragraph_formatting(paragraph, first_line_indent):
    """应用段落格式设置"""
    if not paragraph.style.name.startswith("Heading"):
        p = paragraph._element
        pPr = p.get_or_add_pPr()

        # 设置首行缩进
        ind = pPr.get_or_add_ind()
        ind.set(qn("w:firstLine"), str(int(first_line_indent * 567)))

        # 设置段落间距和行距
        spacing = pPr.get_or_add_spacing()
        spacing.set(qn("w:after"), "100")
        spacing.set(qn("w:line"), "240")
        spacing.set(qn("w:lineRule"), "auto")


def reformat_docx(file_path, output_path=None, first_line_indent=0.74):
    """重新格式化Word文档"""
    doc = Document(file_path)

    if output_path is None:
        output_path = file_path

    for paragraph in doc.paragraphs:
        # 处理换行符
        p = paragraph._element
        br_elements = p.findall(
            ".//w:br",
            {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
        )

        # 处理连续换行符
        handle_consecutive_breaks(br_elements, first_line_indent)

        # 应用段落格式
        apply_paragraph_formatting(paragraph, first_line_indent)

    doc.save(output_path)


if __name__ == "__main__":
    reformat_docx("../results/110_Murong_Jun.docx")
